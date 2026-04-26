from __future__ import annotations

from io import BytesIO
import json

import httpx
from docx import Document as DocxDocument
from openpyxl import Workbook

from scraper.ai.ai_enhancement import AIClassifier, _build_context_prompt_payload, MAX_PROMPT_CHARS
from scraper.fetchers.http_fetcher import HttpFetcher
from scraper.schemas import PageContentDocument, PageFetchResult
from scraper.utils.document_reader import extract_local_document_text, infer_document_kind


def _pdf_bytes(text: str) -> bytes:
    return (
        b"%PDF-1.4\n1 0 obj\n<<>>\nstream\nBT ("
        + text.encode("utf-8")
        + b") Tj ET\nendstream\nendobj\n%%EOF"
    )


def _docx_bytes(text: str) -> bytes:
    buffer = BytesIO()
    document = DocxDocument()
    document.add_paragraph(text)
    document.save(buffer)
    return buffer.getvalue()


def _xlsx_bytes(text: str) -> bytes:
    buffer = BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sheet1"
    sheet["A1"] = text
    workbook.save(buffer)
    return buffer.getvalue()


def _response(url: str, content: bytes, content_type: str) -> httpx.Response:
    request = httpx.Request("GET", url)
    return httpx.Response(200, request=request, content=content, headers={"content-type": content_type})


def test_local_document_text_extraction_supports_pdf_docx_and_xlsx() -> None:
    pdf_result = extract_local_document_text(_pdf_bytes("Green Energy Grant"), "https://example.org/form.pdf", "application/pdf")
    docx_result = extract_local_document_text(
        _docx_bytes("Mentorship and business support"),
        "https://example.org/form.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    xlsx_result = extract_local_document_text(
        _xlsx_bytes("Application deadline"),
        "https://example.org/form.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    assert "Green Energy Grant" in pdf_result.text
    assert "Mentorship and business support" in docx_result.text
    assert "Application deadline" in xlsx_result.text
    assert infer_document_kind("https://example.org/form.png", "image/png") == "image"


def test_http_fetcher_extracts_document_text_without_html(settings, monkeypatch) -> None:
    fetcher = HttpFetcher(settings)
    monkeypatch.setattr(fetcher, "can_fetch", lambda _url: True)

    docx_bytes = _docx_bytes("Business management training")
    monkeypatch.setattr(
        fetcher.client,
        "get",
        lambda url, headers=None: _response(
            url,
            docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
    )

    result = fetcher.fetch("https://example.org/support/application-pack.docx")

    assert result.succeeded is True
    assert "Business management training" in result.html
    assert any("DOCX text extracted" in note for note in result.notes)


def test_image_document_fetch_result_counts_as_success() -> None:
    page = PageFetchResult(
        url="https://example.org/support/poster.png",
        requested_url="https://example.org/support/poster.png",
        canonical_url="https://example.org/support/poster.png",
        final_url="https://example.org/support/poster.png",
        status_code=200,
        content_type="image/png",
        html="",
        title=None,
        fetch_method="http",
        headers={},
        js_rendered=False,
        notes=[],
    )

    assert page.succeeded is True


def test_ai_classifier_reads_document_links_into_prompt_payload(monkeypatch) -> None:
    classifier = AIClassifier(
        {
            "aiProvider": "openai",
            "openaiKey": "test",
            "aiModel": "gpt-test",
            "documentAiMaxDocumentsPerPage": 2,
            "documentAiMaxExtractedChars": 4000,
        },
        storage=None,
    )

    document = PageContentDocument(
        page_url="https://example.org/programmes/green-energy-grant",
        title="Green Energy Grant",
        source_content_type="text/html",
        full_body_text="The grant supports startups.",
        source_domain="example.org",
        document_links=[
            "https://example.org/docs/green-energy-pack.pdf",
            "https://example.org/docs/green-energy-poster.png",
            "https://example.org/docs/tourism-transformation-checklist.pdf",
            "https://other.example.net/ignore-me.pdf",
        ],
    )

    observed_contexts = []

    def fake_read_remote_document(url: str):
        if url.endswith(".pdf"):
            return _pdf_bytes("R250 000 to R1 000 000"), "application/pdf", []
        if url.endswith(".png"):
            return b"fake-image", "image/png", []
        raise AssertionError(f"unexpected document url: {url}")

    def fake_summarize_document_with_openai(**kwargs):
        observed_contexts.append(kwargs.get("programme_context"))
        if kwargs["kind"] == "pdf":
            return {"summary": "Funding between R250 000 and R1 000 000.", "key_points": ["Supports startups"], "notes": []}
        if kwargs["kind"] == "image":
            return {"summary": "Poster for the green energy grant.", "key_points": ["Green Energy Grant"], "notes": []}
        return {}

    monkeypatch.setattr(classifier, "_read_remote_document", fake_read_remote_document)
    monkeypatch.setattr(classifier, "_summarize_document_with_openai", fake_summarize_document_with_openai)

    prepared = classifier._prepare_document_context(document)
    payload = _build_context_prompt_payload(prepared)

    assert len(prepared.document_evidence) == 2
    assert payload["document_evidence"][0]["document_url"] == "https://example.org/docs/green-energy-pack.pdf"
    assert payload["document_evidence_text"]
    assert "Funding between R250 000 and R1 000 000." in payload["document_evidence_text"]
    assert observed_contexts and observed_contexts[0]["programme_name"] == "Green Energy Grant"
    assert len(json.dumps(payload, ensure_ascii=False)) < MAX_PROMPT_CHARS


def test_ai_classifier_merges_document_evidence_into_final_record(tmp_path) -> None:
    classifier = AIClassifier({"disableRemoteAi": True}, storage=None)

    document = PageContentDocument(
        page_url="https://example.org/programmes/green-energy-grant",
        title="Green Energy Grant - National Empowerment Fund",
        source_content_type="application/pdf",
        source_domain="example.org",
        document_evidence=[
            {
                "document_url": "https://example.org/docs/green-energy-pack.pdf",
                "document_kind": "pdf",
                "content_type": "application/pdf",
                "source_method": "page",
                "summary": "Grant offers R250 000 to R1 000 000.",
                "key_points": [
                    "Apply online at https://example.org/apply",
                    "Contact grants@example.org",
                ],
                "extracted_text": "Required documents: ID document and registration certificate.",
            }
        ],
    )

    records = classifier.classify_document(document)

    assert len(records) == 1
    record = records[0]
    assert record.funder_name == "National Empowerment Fund"
    assert record.program_name == "Green Energy Grant"
    assert any("R250 000 to R1 000 000" in item for item in record.raw_documents_data)
    assert any("document evidence read from" in note.lower() for note in record.notes)
    assert record.application_url == "https://example.org/apply"
    assert record.contact_email == "grants@example.org"
    assert any("registration certificate" in item.lower() for item in record.required_documents)
