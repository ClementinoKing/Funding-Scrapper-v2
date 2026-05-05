"""Document extractors for various file formats."""

from __future__ import annotations

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, List, Optional, Union

from scraper.utils.text import clean_text


class DocumentExtractor(ABC):
    """Abstract base for document extractors."""

    @abstractmethod
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """Extract text from document."""
        pass

    @abstractmethod
    def extract_metadata(self, file_path: Union[str, Path]) -> Dict[str, str]:
        """Extract metadata from document."""
        pass

    @abstractmethod
    def supports_file(self, file_path: Union[str, Path]) -> bool:
        """Check if this extractor supports the file."""
        pass


class PDFExtractor(DocumentExtractor):
    """Enhanced PDF text extractor with OCR fallback."""

    def __init__(self, ocr_provider: Optional[object] = None) -> None:
        self.ocr_provider = ocr_provider

    def extract_text(self, file_path: Union[str, Path]) -> str:
        """Extract text from PDF."""
        try:
            import pypdf

            text_parts: List[str] = []
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            extracted_text = "\n".join(text_parts)

            # If text is too short, try OCR
            if len(extracted_text.strip()) < 100 and self.ocr_provider:
                try:
                    ocr_text = self.ocr_provider.extract_text(file_path)
                    if len(ocr_text) > len(extracted_text):
                        extracted_text = ocr_text
                except Exception:
                    pass

            return clean_text(extracted_text)
        except Exception as e:
            # Fallback to OCR if available
            if self.ocr_provider:
                try:
                    return self.ocr_provider.extract_text(file_path)
                except Exception:
                    pass
            return ""

    def extract_metadata(self, file_path: Union[str, Path]) -> Dict[str, str]:
        """Extract metadata from PDF."""
        try:
            import pypdf

            metadata: Dict[str, str] = {}
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                if reader.metadata:
                    for key, value in reader.metadata.items():
                        if value:
                            metadata[key.replace("/", "")] = str(value)
                metadata["page_count"] = str(len(reader.pages))
            return metadata
        except Exception:
            return {}

    def supports_file(self, file_path: Union[str, Path]) -> bool:
        """Check if file is a PDF."""
        path = Path(file_path)
        return path.suffix.lower() == ".pdf"


class WordExtractor(DocumentExtractor):
    """Word document (.docx) text extractor."""

    def extract_text(self, file_path: Union[str, Path]) -> str:
        """Extract text from Word document."""
        try:
            import docx

            doc = docx.Document(file_path)
            text_parts: List[str] = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text:
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text:
                        text_parts.append(row_text)

            return clean_text("\n".join(text_parts))
        except Exception:
            return ""

    def extract_metadata(self, file_path: Union[str, Path]) -> Dict[str, str]:
        """Extract metadata from Word document."""
        try:
            import docx

            doc = docx.Document(file_path)
            metadata: Dict[str, str] = {}

            if doc.core_properties:
                props = doc.core_properties
                if props.title:
                    metadata["title"] = props.title
                if props.author:
                    metadata["author"] = props.author
                if props.created:
                    metadata["created"] = str(props.created)
                if props.modified:
                    metadata["modified"] = str(props.modified)

            return metadata
        except Exception:
            return {}

    def supports_file(self, file_path: Union[str, Path]) -> bool:
        """Check if file is a Word document."""
        path = Path(file_path)
        return path.suffix.lower() in {".docx", ".doc"}


class ExcelExtractor(DocumentExtractor):
    """Excel spreadsheet (.xlsx) text extractor."""

    def extract_text(self, file_path: Union[str, Path]) -> str:
        """Extract text from Excel spreadsheet."""
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_parts: List[str] = []

            for sheet in workbook.worksheets:
                text_parts.append(f"Sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        text_parts.append(row_text)

            return clean_text("\n".join(text_parts))
        except Exception:
            return ""

    def extract_metadata(self, file_path: Union[str, Path]) -> Dict[str, str]:
        """Extract metadata from Excel spreadsheet."""
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(file_path, data_only=True)
            metadata: Dict[str, str] = {}

            if workbook.properties:
                props = workbook.properties
                if props.title:
                    metadata["title"] = props.title
                if props.creator:
                    metadata["creator"] = props.creator
                if props.created:
                    metadata["created"] = str(props.created)
                if props.modified:
                    metadata["modified"] = str(props.modified)

            metadata["sheet_count"] = str(len(workbook.worksheets))
            metadata["sheet_names"] = ", ".join(sheet.title for sheet in workbook.worksheets)

            return metadata
        except Exception:
            return {}

    def supports_file(self, file_path: Union[str, Path]) -> bool:
        """Check if file is an Excel spreadsheet."""
        path = Path(file_path)
        return path.suffix.lower() in {".xlsx", ".xls"}


def get_extractor_for_file(
    file_path: Union[str, Path],
    ocr_provider: Optional[object] = None,
) -> Optional[DocumentExtractor]:
    """Get appropriate extractor for a file."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return PDFExtractor(ocr_provider=ocr_provider)
    elif suffix in {".docx", ".doc"}:
        return WordExtractor()
    elif suffix in {".xlsx", ".xls"}:
        return ExcelExtractor()

    return None
