"""Helpers for reading and compacting document content."""

from __future__ import annotations

import io
from dataclasses import dataclass
from typing import List, Optional, Tuple

from scraper.utils.pdf import extract_pdf_text
from scraper.utils.text import clean_text, unique_preserve_order

try:  # pragma: no cover - optional dependency may be missing in lean environments
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover - handled gracefully at runtime
    DocxDocument = None  # type: ignore[assignment]

try:  # pragma: no cover - optional dependency may be missing in lean environments
    from openpyxl import load_workbook
except Exception:  # pragma: no cover - handled gracefully at runtime
    load_workbook = None  # type: ignore[assignment]


DOCUMENT_KIND_PDF = "pdf"
DOCUMENT_KIND_DOCX = "docx"
DOCUMENT_KIND_XLSX = "xlsx"
DOCUMENT_KIND_IMAGE = "image"
DOCUMENT_KIND_TEXT = "text"
DOCUMENT_KIND_UNSUPPORTED = "unsupported"

DOCUMENT_TEXT_EXTENSIONS = (".pdf", ".doc", ".docx", ".xls", ".xlsx")
DOCUMENT_IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tif", ".tiff")
DOCUMENT_BINARY_SKIP_CONTENT_TYPES = (
    "application/zip",
    "application/x-zip-compressed",
    "application/x-rar-compressed",
    "application/x-7z-compressed",
    "application/octet-stream",
    "application/x-msdownload",
)

PDF_CONTENT_TYPES = (
    "application/pdf",
    "application/x-pdf",
)
DOCX_CONTENT_TYPES = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
XLSX_CONTENT_TYPES = (
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
)
IMAGE_CONTENT_PREFIXES = ("image/",)
TEXT_CONTENT_PREFIXES = ("text/", "application/xml", "application/xhtml+xml")


@dataclass(frozen=True)
class DocumentTextResult:
    kind: str
    text: str
    notes: Tuple[str, ...] = ()


def infer_document_kind(url: str, content_type: Optional[str] = None) -> str:
    lowered_url = (url or "").lower()
    lowered_content_type = (content_type or "").lower()

    if lowered_content_type.startswith(IMAGE_CONTENT_PREFIXES) or any(lowered_url.endswith(ext) for ext in DOCUMENT_IMAGE_EXTENSIONS):
        return DOCUMENT_KIND_IMAGE
    if lowered_content_type.startswith(PDF_CONTENT_TYPES) or lowered_url.endswith(".pdf"):
        return DOCUMENT_KIND_PDF
    if lowered_content_type.startswith(DOCX_CONTENT_TYPES) or lowered_url.endswith(".docx"):
        return DOCUMENT_KIND_DOCX
    if lowered_content_type.startswith(XLSX_CONTENT_TYPES) or lowered_url.endswith(".xlsx"):
        return DOCUMENT_KIND_XLSX
    if lowered_content_type.startswith(TEXT_CONTENT_PREFIXES):
        return DOCUMENT_KIND_TEXT
    if lowered_url.endswith(".doc") or lowered_url.endswith(".xls") or lowered_url.endswith(".ppt") or lowered_url.endswith(".pptx"):
        return DOCUMENT_KIND_UNSUPPORTED
    return DOCUMENT_KIND_UNSUPPORTED


def is_skip_content_type(content_type: Optional[str], skip_content_types: Sequence[str]) -> bool:
    lowered = (content_type or "").lower()
    return any(lowered.startswith(pattern.lower()) for pattern in skip_content_types)


def _extract_docx_text(document_bytes: bytes, *, max_parts: int = 200) -> str:
    if DocxDocument is None:
        return ""
    try:
        doc = DocxDocument(io.BytesIO(document_bytes))
    except Exception:
        return ""
    chunks: List[str] = []
    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            chunks.append(text)
        if len(chunks) >= max_parts:
            break
    if len(chunks) < max_parts:
        for table in doc.tables:
            for row in table.rows:
                cells = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
                if cells:
                    chunks.append(" | ".join(cells))
                if len(chunks) >= max_parts:
                    break
            if len(chunks) >= max_parts:
                break
    return "\n".join(unique_preserve_order(chunks)).strip()


def _extract_xlsx_text(document_bytes: bytes, *, max_rows: int = 200, max_cells: int = 2000) -> str:
    if load_workbook is None:
        return ""
    try:
        workbook = load_workbook(io.BytesIO(document_bytes), read_only=True, data_only=True)
    except Exception:
        return ""

    chunks: List[str] = []
    cell_count = 0
    try:
        for sheet in workbook.worksheets:
            sheet_chunks: List[str] = []
            for row in sheet.iter_rows(values_only=True):
                values = [clean_text(str(cell)) for cell in row if cell is not None and clean_text(str(cell))]
                if values:
                    sheet_chunks.append("\t".join(values))
                    cell_count += len(values)
                if len(sheet_chunks) >= max_rows or cell_count >= max_cells:
                    break
            if sheet_chunks:
                chunks.append(f"[Sheet: {sheet.title}]")
                chunks.extend(sheet_chunks)
            if cell_count >= max_cells:
                break
    finally:
        try:
            workbook.close()
        except Exception:
            pass

    return "\n".join(unique_preserve_order(chunks)).strip()


def extract_local_document_text(document_bytes: bytes, url: str, content_type: Optional[str] = None) -> DocumentTextResult:
    """Extract text from a document when possible, with a stable kind label."""

    kind = infer_document_kind(url, content_type)
    notes: List[str] = []
    if not document_bytes:
        return DocumentTextResult(kind=kind, text="", notes=("empty document body",))

    if kind == DOCUMENT_KIND_PDF:
        text = extract_pdf_text(document_bytes)
        if not text:
            notes.append("pdf text extraction yielded no readable text")
        return DocumentTextResult(kind=kind, text=text, notes=tuple(notes))

    if kind == DOCUMENT_KIND_DOCX:
        text = _extract_docx_text(document_bytes)
        if not text:
            notes.append("docx text extraction yielded no readable text")
        return DocumentTextResult(kind=kind, text=text, notes=tuple(notes))

    if kind == DOCUMENT_KIND_XLSX:
        text = _extract_xlsx_text(document_bytes)
        if not text:
            notes.append("xlsx text extraction yielded no readable text")
        return DocumentTextResult(kind=kind, text=text, notes=tuple(notes))

    if kind == DOCUMENT_KIND_IMAGE:
        notes.append("image document requires OpenAI reading")
        return DocumentTextResult(kind=kind, text="", notes=tuple(notes))

    if kind == DOCUMENT_KIND_TEXT:
        text = clean_text(document_bytes.decode("utf-8", errors="ignore"))
        return DocumentTextResult(kind=kind, text=text, notes=tuple(notes))

    return DocumentTextResult(kind=kind, text="", notes=("unsupported document type",))


def compact_document_text(text: str, *, max_chars: int) -> str:
    cleaned = clean_text(text or "")
    if not cleaned:
        return ""
    normalized = "\n".join(unique_preserve_order([line for line in cleaned.splitlines() if line.strip()]))
    return normalized[:max_chars]


def document_type_label(url: str, content_type: Optional[str] = None) -> str:
    return infer_document_kind(url, content_type)
